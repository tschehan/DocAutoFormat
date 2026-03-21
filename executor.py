"""Minimal plan executor."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from schemas import ExecutionPlan, PlanAction

FONT_SIZE_MAP = {
    "二号": Pt(22),
    "三号": Pt(16),
    "四号": Pt(14),
    "小四": Pt(12),
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

LINE_SPACING_MAP = {
    "1": 1.0,
    "1.0": 1.0,
    "1.5": 1.5,
    "2": 2.0,
    "2.0": 2.0,
}


def _parse_paragraph_index(block_id: str) -> int | None:
    """Get paragraph index from a numeric block_id like 0."""
    if not block_id.isdigit():
        return None
    return int(block_id)


def _set_rpr_font_names(
    run,
    east_asia_font_name: str | None,
    western_font_name: str | None,
) -> None:
    """Write Chinese and Western font names separately."""
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()

    if east_asia_font_name is not None:
        r_fonts.set(qn("w:eastAsia"), east_asia_font_name)
    if western_font_name is not None:
        run.font.name = western_font_name
        r_fonts.set(qn("w:ascii"), western_font_name)
        r_fonts.set(qn("w:hAnsi"), western_font_name)
        r_fonts.set(qn("w:cs"), western_font_name)

    for attr_name in (
        "w:asciiTheme",
        "w:hAnsiTheme",
        "w:eastAsiaTheme",
        "w:cstheme",
    ):
        attr_key = qn(attr_name)
        if attr_key in r_fonts.attrib:
            del r_fonts.attrib[attr_key]


def _set_rpr_font_size(run, font_size) -> None:
    """Write font size to both normal and complex-script nodes."""
    run.font.size = font_size
    half_points = str(int(font_size.pt * 2))
    r_pr = run._element.get_or_add_rPr()

    for tag_name in ("w:sz", "w:szCs"):
        size_node = r_pr.find(qn(tag_name))
        if size_node is None:
            size_node = OxmlElement(tag_name)
            r_pr.append(size_node)
        size_node.set(qn("w:val"), half_points)


def _apply_run_style(paragraph, action: PlanAction) -> None:
    """Apply run-level settings."""
    font_size = FONT_SIZE_MAP.get(action.font_size) if action.font_size else None

    for run in paragraph.runs:
        if action.font_name is not None or action.western_font_name is not None:
            _set_rpr_font_names(run, action.font_name, action.western_font_name)
        if font_size is not None:
            _set_rpr_font_size(run, font_size)
        if action.bold is not None:
            run.font.bold = action.bold


def _apply_paragraph_style(paragraph, action: PlanAction) -> None:
    """Apply paragraph-level settings."""
    if action.line_spacing is not None:
        line_spacing = LINE_SPACING_MAP.get(action.line_spacing)
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing

    if action.alignment is not None:
        alignment = ALIGNMENT_MAP.get(action.alignment)
        if alignment is not None:
            paragraph.alignment = alignment


def _describe_first_run(paragraph) -> tuple[str | None, float | None, bool | None]:
    """Read the first run after applying styles."""
    if not paragraph.runs:
        return None, None, None

    first_run = paragraph.runs[0]
    font_name = first_run.font.name
    size_pt = first_run.font.size.pt if first_run.font.size is not None else None
    bold = first_run.font.bold
    return font_name, size_pt, bold


def _describe_rfonts(paragraph) -> str:
    """Read raw rFonts attrs from the first run."""
    if not paragraph.runs:
        return "no-runs"

    first_run = paragraph.runs[0]
    r_pr = first_run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return "no-rFonts"

    r_fonts = r_pr.rFonts
    parts = []
    for attr_name in (
        "w:ascii",
        "w:hAnsi",
        "w:eastAsia",
        "w:cs",
        "w:asciiTheme",
        "w:hAnsiTheme",
        "w:eastAsiaTheme",
        "w:cstheme",
    ):
        value = r_fonts.get(qn(attr_name))
        if value is not None:
            parts.append(f"{attr_name}={value}")
    return ", ".join(parts) if parts else "empty-rFonts"


def execute_plan(
    input_file: str | ExecutionPlan,
    output_file: str | None = None,
    plan: ExecutionPlan | None = None,
) -> dict:
    """Execute a plan and keep old compatibility."""
    if isinstance(input_file, ExecutionPlan) and output_file is None and plan is None:
        return {
            "status": "pending",
            "applied_action_count": 0,
            "skipped_action_count": len(input_file.actions),
        }

    if not isinstance(input_file, str) or output_file is None or plan is None:
        raise ValueError("Use execute_plan(input_file, output_file, plan).")

    doc = Document(input_file)
    applied_action_count = 0
    skipped_action_count = 0

    for action in plan.actions:
        paragraph_index = _parse_paragraph_index(action.target_block_id)
        if paragraph_index is None:
            print(f"[SKIP] {action.target_block_id} | unsupported block type")
            skipped_action_count += 1
            continue

        if paragraph_index >= len(doc.paragraphs):
            print(f"[SKIP] {action.target_block_id} | paragraph out of range")
            skipped_action_count += 1
            continue

        print(
            f"[EXEC] {action.target_block_id} | role={action.role} | "
            f"zh_font={action.font_name} | en_font={action.western_font_name} | "
            f"size={action.font_size} | bold={action.bold} | "
            f"spacing={action.line_spacing} | align={action.alignment}"
        )

        paragraph = doc.paragraphs[paragraph_index]
        _apply_run_style(paragraph, action)
        _apply_paragraph_style(paragraph, action)

        font_name, size_pt, bold = _describe_first_run(paragraph)
        alignment = paragraph.alignment.name if paragraph.alignment is not None else None
        line_spacing = paragraph.paragraph_format.line_spacing
        rfonts = _describe_rfonts(paragraph)
        print(
            f"[APPLIED] {action.target_block_id} | font={font_name} | "
            f"size_pt={size_pt} | bold={bold} | alignment={alignment} | "
            f"line_spacing={line_spacing} | rFonts={rfonts}"
        )

        applied_action_count += 1

    doc.save(output_file)

    return {
        "status": "done",
        "applied_action_count": applied_action_count,
        "skipped_action_count": skipped_action_count,
    }
