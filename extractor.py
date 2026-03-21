"""Minimal document extractor."""

from pathlib import Path

from docx import Document

from schemas import DocumentBlock, StructureResult


def extract_document_info(file_path: str) -> StructureResult:
    """Read a .docx file and extract paragraphs and table cells."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if path.suffix.lower() != ".docx":
        raise ValueError("Only .docx files are supported.")

    doc = Document(file_path)
    blocks: list[DocumentBlock] = []

    for index, paragraph in enumerate(doc.paragraphs):
        blocks.append(
            DocumentBlock(
                block_id=str(index),
                text=paragraph.text,
                role=None,
            )
        )

    for table_index, table in enumerate(doc.tables):
        table_id = f"table_{table_index}"
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                blocks.append(
                    DocumentBlock(
                        block_id=f"t_{table_index}_{row_index}_{col_index}",
                        text=cell.text,
                        role="table_cell",
                        table_id=table_id,
                        row=row_index,
                        col=col_index,
                    )
                )

    return StructureResult(blocks=blocks)


def print_structure(result: StructureResult) -> None:
    """Print a short preview of extracted blocks."""
    print(f"Total blocks: {len(result.blocks)}")

    for block in result.blocks[:5]:
        preview = block.text.replace("\n", " ").strip()
        if len(preview) > 40:
            preview = preview[:40] + "..."
        print(f"{block.block_id} | {preview}")
